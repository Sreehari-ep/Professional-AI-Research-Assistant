from docx import Document
from reportlab.platypus import SimpleDocTemplate,Paragraph,Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from utils.paths import REPORTS_FOLDER
def export_txt(title,sections,filename):
    path=REPORTS_FOLDER/f"{filename}.txt"; path.write_text(title+"\n\n"+"\n\n".join(f"{k}\n{v}" for k,v in sections.items()),encoding="utf-8"); return path
def export_docx(title,sections,filename):
    path=REPORTS_FOLDER/f"{filename}.docx"; doc=Document(); doc.add_heading(title,0)
    for k,v in sections.items(): doc.add_heading(k,1); doc.add_paragraph(str(v))
    doc.save(path); return path
def export_pdf(title,sections,filename):
    path=REPORTS_FOLDER/f"{filename}.pdf"; styles=getSampleStyleSheet(); story=[Paragraph(title,styles["Title"]),Spacer(1,12)]
    for k,v in sections.items(): story += [Paragraph(k,styles["Heading1"]),Paragraph(str(v).replace("\n","<br/>"),styles["BodyText"]),Spacer(1,8)]
    SimpleDocTemplate(str(path),pagesize=A4).build(story); return path
